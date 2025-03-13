import os
import ebooklib # Parse EPUB
import pymupdf # Parse PDF
import math 

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from ebooklib import epub
from bs4 import BeautifulSoup # Parse HTML, XML


def html_to_str(chapter):
    soup = BeautifulSoup(chapter.get_content(), 'html.parser') # Built-in Parser
    return soup.get_text()

def clean_str(str):
    str = str.replace('\xa0', ' ').replace('\n', ' ').replace('  ', ' ') # Deal with non-text content. 
    return str.strip()

# Chapter-based chunks. Iterate through pages in pdf until we get to the hyperlinks (table of contents) to get the pages corresponding to chapters
# Assumes that the first instance of hyperlinks is the table of contents
# If there are no hyperlinks just returns [start, end] pages. 
def extract_pages_hyperlinks_pdf(path):
    doc = pymupdf.open(path)
    hyperlinks = []
    hyperlinks.append(0) 
    for page in doc:
        links = page.get_links()  # Extract all links
        need_to_sort = False
        if len(links) > 0: # Table of Contents
            for link in links:
                temp = link.get("page") 
                if temp not in hyperlinks:
                    if len(hyperlinks) > 0:
                        if temp < hyperlinks[-1]:
                            need_to_sort = True
                        if (temp - 2) < hyperlinks[-1]: # each page is ~500 tokens max. 
                            continue
                    hyperlinks.append(temp)
            hyperlinks.append(len(doc))
            if need_to_sort:
                sorted(hyperlinks) # When the hyperlinks are out of order
            return doc, hyperlinks

    hyperlinks.append(len(doc))
    doc, hyperlinks

def parse_epub(path): 
    book = epub.read_epub(path)
    items = list(book.get_items_of_type(ebooklib.ITEM_DOCUMENT)) # ITEM_DOCUMENT is flag for actual words
    chapters = []
    for item in items:
        temp = clean_str(html_to_str(item))
        # temp will be empty if it's the item is non-content data
        if temp:
            chapters.append(temp)
    return chapters

def parse_pdf(path): 
    # Create Chunks Based on "chapters" via Hyperlink pages. 
    # If Hyperlinks don't exist
    doc, hyperlinks = extract_pages_hyperlinks_pdf(path)
    chapters = []
    for i in range(0, len(hyperlinks) - 1):
        text = []
        for pages in range(hyperlinks[i], hyperlinks[i + 1]):
            text.append(clean_str(doc[pages].get_text()))
        chapters.append(''.join(text))
    if len(hyperlinks) == 2:
        chapters = generate_equal_chunks(chapters)
    return chapters

def parse_xml(path):
    soup = BeautifulSoup(open(path, "r", encoding="utf-8").read(), 'lxml') # Built-in fast xml parser 
    title_tag = soup.find("title") # Add the title to the content
    chapters = []

    if title_tag:
        chapters.append(title_tag.text.strip()) # clean 
    content_tags = soup.find_all("content:encoded") # all actual text is in here
    for content in content_tags: 
        temp = content.text
        if temp.endswith("]]>"):
            temp = temp[:-3] # extra non-text content to be removed
        temp = clean_str(temp) # final clean
        if temp:
            chapters.append(temp)
    return chapters

def parser(path): # singular call for all files
    file, extension = os.path.splitext(path)
    if extension == ".epub":
        return parse_epub(path)
    elif extension == ".pdf":
        return parse_pdf(path)
    elif extension == '.xml':
        return parse_xml(path)
    else: 
        file = open(path, "r", encoding = "utf-8")
        return file.read()

def coalesce(book, min_length=15000): # make the chunks more "even"
    out = []
    for str in book:
        # Combing smaller chapters or extra non-chapter information into same call
        if len(out) > 0 and (len(out[-1]) + len(str) < min_length or len(str) > 10 * len(out[-1])):
            temp = out[-1]
            out[-1] = temp + " " + str
        else:
            out.append(str)
    return out

def generate_equal_chunks(coalesed_books, num_calls = 10): # make the chunks all equal to each other (as best as possible)
    output = []
    for book in coalesed_books:
        min_size = 0
        for chapter in book:
            min_size += len(chapter)
        min_size = math.ceil(float(min_size) / num_calls) # num_calls of each min_size characters (tokens ~ characters / 4)

        combined = []
        current = ""

        for chapter in book:
            if len(current) + len(chapter) <= min_size: # If adding this chapter is still under min_size just add it and continue
                current = current + " " + chapter
            else: 
                combined_string = current + " " + chapter
                split = min_size
                punctuation = {'.', '!', '?'}
                for chr in range(min_size - 1, -1, -1): # Find the index of last sentence before this min_size index.  
                    if combined_string[chr] in punctuation:
                        split = chr + 1
                        break

                combined.append(combined_string[:split].strip()) #
                current = combined_string[split:] # First Sentence in currrent chapter not in previous iteration.
        if current: 
            combined.append(current.strip()) 
        if combined: 
            output.append(combined)
    return output


def read_doc(path): # Read .docx files, useful for debugging
  doc = Document(path)
  return "\n".join([p.text for p in doc.paragraphs])

def save_to_doc(report, filename): # Final form of output 
    doc = Document()
    paragraph = report.split('\n')
    heading = doc.add_heading("Comparative Analysis", level=1) 
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for text in paragraph:
      if not text: 
        continue
      doc.add_paragraph("       " + text) # indent of paragraphs in 5-paragraph essay
    doc.save(filename)
